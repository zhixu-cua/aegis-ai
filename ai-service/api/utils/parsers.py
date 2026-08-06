import os
import warnings
import threading

from paddleocr import PaddleOCR

def parse_txt(file_path: str) -> str:
    try:
        with open(file_path, "r", encoding="utf-8") as f:
            return f.read()
    except UnicodeDecodeError:
        with open(file_path, "r", encoding="gbk") as f:
            return f.read()

def parse_pdf(file_path: str) -> str:
    try:
        import pdfplumber
    except ImportError:
        raise Exception("缺少 pdfplumber 库，请执行 pip install pdfplumber")
        
    text = ""
    with pdfplumber.open(file_path) as pdf:
        for page_idx, page in enumerate(pdf.pages):
            page_text = page.extract_text()
            
            # 如果当前页提取不到文本，或者文本极少（<20字符），判定为扫描件或纯图片页
            if not page_text or len(page_text.strip()) < 20:
                import os
                import tempfile
                try:
                    import fitz  # PyMuPDF
                except ImportError:
                    raise Exception("检测到扫描版 PDF，需要 PyMuPDF 库将其转为图片，请执行 pip install PyMuPDF")
                
                doc = fitz.open(file_path)
                fitz_page = doc.load_page(page_idx)
                # 提升分辨率以增加 OCR 准确率（2倍缩放）
                mat = fitz.Matrix(2.0, 2.0)
                pix = fitz_page.get_pixmap(matrix=mat)
                
                temp_dir = tempfile.gettempdir()
                temp_img_path = os.path.join(temp_dir, f"temp_pdf_page_{page_idx}.png")
                pix.save(temp_img_path)
                doc.close()
                
                try:
                    # 调用下方已有的 parse_image 函数对图片进行 OCR 识别
                    ocr_text = parse_image(temp_img_path)
                    if ocr_text:
                        text += ocr_text + "\n\n"
                finally:
                    if os.path.exists(temp_img_path):
                        try:
                            os.remove(temp_img_path)
                        except:
                            pass
            else:
                text += page_text + "\n\n"
                
            # 尝试提取表格 (若是纯图片扫描件，此处通常不会报错，只是提取为空)
            tables = page.extract_tables()
            for table in tables:
                if not table: continue
                for row in table:
                    clean_row = [str(cell).replace('\n', ' ') if cell else "" for cell in row]
                    text += "| " + " | ".join(clean_row) + " |\n"
                text += "\n"
    return text

def parse_docx(file_path: str) -> str:
    try:
        from docx import Document
        doc = Document(file_path)
        text = ""
        for para in doc.paragraphs:
            if para.text.strip():
                text += para.text + "\n"
        for table in doc.tables:
            for row in table.rows:
                row_data = [cell.text.replace('\n', ' ') for cell in row.cells]
                text += "| " + " | ".join(row_data) + " |\n"
            text += "\n"
        return text
    except ImportError:
        raise Exception("缺少 python-docx 库，请执行 pip install python-docx")

def parse_doc(file_path: str) -> str:
    """
    处理旧版 .doc 文件：通过 win32com 调用 Windows 系统的 Word 将其转换为临时 .docx，
    然后复用 parse_docx 逻辑，从而完美保留原有的段落和表格 Markdown 结构。
    """
    import os
    import tempfile
    try:
        import win32com.client
        import pythoncom
    except ImportError:
        raise Exception("缺少 pywin32 库，请执行 pip install pywin32")

    # 在 Redis 后台线程/异步任务中调用 COM 接口，必须先初始化
    pythoncom.CoInitialize()
    
    abs_file_path = os.path.abspath(file_path)
    temp_dir = tempfile.gettempdir()
    temp_docx_path = os.path.join(temp_dir, os.path.basename(file_path) + "x")

    try:
        word = win32com.client.DispatchEx("Word.Application")
        word.Visible = False
        word.DisplayAlerts = 0  # wdAlertsNone: 关闭所有弹窗警告
        # 打开 .doc 文件 (设置只读，禁用转换弹窗，绕过部分保护模式拦截)
        doc = word.Documents.Open(abs_file_path, ReadOnly=True, ConfirmConversions=False)
        # 16 代表 wdFormatXMLDocument (.docx)
        doc.SaveAs(temp_docx_path, 16)
        doc.Close(False)
        word.Quit()
        
        # 复用 parse_docx 以提取文本和表格 Markdown
        return parse_docx(temp_docx_path)
    except Exception as e:
        raise Exception(f"解析 .doc 失败（需确保本机已安装 Microsoft Word，且文件未被 Office 信任中心拦截）。详细错误: {str(e)}")
    finally:
        # 清理临时转换的 docx 文件
        if os.path.exists(temp_docx_path):
            try:
                os.remove(temp_docx_path)
            except:
                pass
        pythoncom.CoUninitialize()

def parse_excel(file_path: str) -> str:
    try:
        import pandas as pd
        # Read all sheets, works for both .xlsx and .xls (if xlrd is installed)
        dfs = pd.read_excel(file_path, sheet_name=None)
        text = ""
        for sheet_name, df in dfs.items():
            text += f"### 表格：{sheet_name}\n"
            text += df.to_markdown(index=False) + "\n\n"
        return text
    except ImportError as e:
        if "xlrd" in str(e):
            raise Exception("缺少 xlrd 库，无法解析老版 xls 文件。请执行 pip install xlrd")
        if "tabulate" in str(e):
            raise Exception("缺少 tabulate 库，无法将表格转为 Markdown。请执行 pip install tabulate")
        raise Exception(f"缺少解析 Excel 所需的依赖库，请执行 pip install pandas openpyxl tabulate xlrd。详细错误: {str(e)}")

def parse_csv(file_path: str) -> str:
    try:
        import pandas as pd
        import os
        # 尝试使用不同的编码读取 CSV
        try:
            df = pd.read_csv(file_path, encoding='utf-8')
        except UnicodeDecodeError:
            df = pd.read_csv(file_path, encoding='gbk')
            
        text = f"### 表格：{os.path.basename(file_path)}\n"
        text += df.to_markdown(index=False) + "\n\n"
        return text
    except ImportError as e:
        if "tabulate" in str(e):
            raise Exception("缺少 tabulate 库，无法将表格转为 Markdown。请执行 pip install tabulate")
        raise Exception(f"缺少解析 CSV 所需的依赖库，请执行 pip install pandas tabulate。详细错误: {str(e)}")
    except Exception as e:
        raise Exception(f"解析 CSV 失败: {str(e)}")

# ---------- 全局 OCR 单例（线程安全） ----------
_ocr_lock = threading.Lock()
_ocr_instance = None

def _get_ocr(use_angle_cls: bool = True, lang: str = "ch"):
    """惰性创建单例 OCR 对象，复用模型减少加载时间"""
    global _ocr_instance
    if _ocr_instance is None:
        with _ocr_lock:
            if _ocr_instance is None:
                # 设置必要环境变量（适配 2.6.2 版本）
                os.environ.setdefault('FLAGS_enable_pir_api', '0')   # 关闭新 IR，避免潜在问题
                os.environ.setdefault('FLAGS_use_mkldnn', '0')       # 禁用 OneDNN（CPU 加速库）
                os.environ.setdefault('FLAGS_use_onednn', '0')
                os.environ.setdefault('GLOG_v', '0')                 # 只显示错误日志
                
                _ocr_instance = PaddleOCR(
                    use_angle_cls=use_angle_cls,
                    lang=lang,
                    use_gpu=False,          # 若需 GPU 改为 True
                    show_log=False          # 关闭内部冗长日志
                )
    return _ocr_instance

def parse_image(
    file_path: str,
    lang: str = "ch",
    use_angle_cls: bool = True,
    confidence_threshold: float = 0.5,
) -> str:
    try:
        ocr = _get_ocr(use_angle_cls=use_angle_cls, lang=lang)
    except ImportError:
        raise Exception("缺少 paddleocr 库，请执行 pip install paddlepaddle paddleocr")
    except Exception as e:
        raise Exception(f"初始化 PaddleOCR 失败: {str(e)}")

    try:
        result = ocr.ocr(file_path)   # 注意：不传 cls 参数
    except Exception as e:
        raise Exception(f"OCR 识别过程出错: {str(e)}")

    if not result or not result[0]:
        return ""

    text = ""
    for idx, res in enumerate(result):
        if not res:
            continue
        for line in res:
            box = line[0]
            line_text, confidence = line[1]
            if confidence < confidence_threshold:
                continue
            text += line_text + "\n"
        print(f"文本: {text}")
    return text.strip()

def parse_document_content(file_path: str) -> str:
    ext = file_path.split('.')[-1].lower()
    
    if ext in ['txt', 'md']:
        return parse_txt(file_path)
    elif ext == 'pdf':
        return parse_pdf(file_path)
    elif ext == 'docx':
        return parse_docx(file_path)
    elif ext == 'doc':
        return parse_doc(file_path)
    elif ext in ['xls', 'xlsx']:
        return parse_excel(file_path)
    elif ext == 'csv':
        return parse_csv(file_path)
    elif ext in ['png', 'jpg', 'jpeg']:
        return parse_image(file_path)
    else:
        raise ValueError(f"不支持的文件格式: {ext}")

